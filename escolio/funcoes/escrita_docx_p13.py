"""Escrita real de comentário do Word a partir de `P13Comment` — o gap mais
concreto encontrado na sessão de 2026-08-14: até aqui, `P13Comment` e
`InterventionRecord` eram dataclasses puras, sem nenhum I/O sobre um
arquivo `.docx` real. `python-docx` 1.2.0 (já uma dependência do projeto,
usada para leitura em `escolio/ingestao/parser_docx.py`) tem `add_comment`
nativo desde 1.1.0 — não é preciso manipular XML na mão.

## Escopo desta sessão

Só `P13Comment.unit_id` que resolve para um `Paragrafo` do corpo do
documento (prefixo `PAR-`, produzido por `parse_docx`). Citação recuada
(`CIT-`) e nota de rodapé (`NOTA-`/footnote, parte XML separada do corpo)
não são endereçáveis pelo `add_comment` de parágrafo simples — comentário
cujo `unit_id` não resolve fica em `nao_aplicados`, nunca é descartado
silenciosamente [CLAUDE.md §11].

## Nunca sobrescreve o original

`escrever_comentarios_no_docx` recusa `caminho_saida == caminho_docx` —
preservação do original é precondição do teto de intervenção do P13
(`PROPOSTA`/INT-05, nunca aplicado no lugar) [CLAUDE.md §6, P06/02].
"""

from __future__ import annotations

from dataclasses import dataclass, field

import docx
from docx.document import Document as DocumentoDocx
from docx.text.paragraph import Paragraph

from escolio.comentarios.comentario import P13Comment
from escolio.ingestao.parser_docx import _e_citacao_recuada, _paragrafo_e_titulo, parse_docx


@dataclass
class ResultadoDeEscritaDocx:
    caminho_saida: str
    comentarios_aplicados: list[str] = field(default_factory=list)
    nao_aplicados: dict[str, str] = field(default_factory=dict)
    """`comment_id` -> motivo. Nunca silenciado — todo `P13Comment` que não
    vira comentário real no arquivo aparece aqui com o porquê."""


def _mapear_unit_id_para_paragrafo(
    documento_docx: DocumentoDocx, caminho_docx: str
) -> dict[str, Paragraph]:
    """Reusa a MESMA lógica de ordinal de `parse_docx` (skip de parágrafo
    vazio, título não incrementa o ordinal, citação recuada incrementa mas
    não vira `Paragrafo`) para religar `unit_id` ao objeto `Paragraph` real
    — `unit_id` não é um índice direto em `documento_docx.paragraphs`
    [achado desta sessão, ver docstring do módulo]. As duas travessias
    (aqui e em `parse_docx`) precisam do MESMO arquivo de origem; usar
    `documento_docx` já aberto evita abrir o arquivo duas vezes com
    instâncias diferentes de `Document`, que `add_comment`/`save` não
    podem misturar."""
    documento_ingerido = parse_docx(caminho_docx)
    ordinal_por_unit_id = {p.unit_id: p.paragrafo_ordinal for p in documento_ingerido.paragrafos}

    paragrafo_por_ordinal: dict[int, Paragraph] = {}
    ordinal_corpo = 0
    for paragrafo in documento_docx.paragraphs:
        if not paragrafo.text.strip():
            continue
        if _paragrafo_e_titulo(paragrafo):
            continue
        if _e_citacao_recuada(paragrafo):
            ordinal_corpo += 1
            continue
        paragrafo_por_ordinal[ordinal_corpo] = paragrafo
        ordinal_corpo += 1

    return {
        unit_id: paragrafo_por_ordinal[ordinal]
        for unit_id, ordinal in ordinal_por_unit_id.items()
        if ordinal in paragrafo_por_ordinal
    }


def _texto_do_comentario(c: P13Comment) -> str:
    """Corpo do comentário real, a partir dos campos de `P13Comment`
    [§31.5] — nunca reescreve o texto do autor [teto INT-05, §4.4]."""
    partes = [c.problem, f"Evidência: {c.evidence}", f"Impacto: {c.impact}", f"Recomendação: {c.recommended_action}"]
    return "\n\n".join(p for p in partes if p)


def escrever_comentarios_no_docx(
    caminho_docx: str,
    comentarios: list[P13Comment],
    caminho_saida: str,
    *,
    autor: str = "Escólio (sistema) — rascunho não homologado",
    iniciais: str | None = "ESC",
) -> ResultadoDeEscritaDocx:
    if caminho_saida == caminho_docx:
        raise ValueError(
            "caminho_saida não pode ser igual a caminho_docx — o original nunca é sobrescrito "
            "[CLAUDE.md §6, preservação do original]"
        )

    documento_docx = docx.Document(caminho_docx)
    mapa = _mapear_unit_id_para_paragrafo(documento_docx, caminho_docx)

    resultado = ResultadoDeEscritaDocx(caminho_saida=caminho_saida)
    for c in comentarios:
        paragrafo = mapa.get(c.unit_id)
        if paragrafo is None:
            resultado.nao_aplicados[c.comment_id] = (
                f"unit_id={c.unit_id!r} não resolve para um Paragrafo do corpo — citação recuada, "
                "nota de rodapé ou tabela não são endereçáveis por add_comment de parágrafo nesta "
                "sessão [escrita_docx_p13.py, escopo declarado]"
            )
            continue
        if not paragrafo.runs:
            resultado.nao_aplicados[c.comment_id] = f"unit_id={c.unit_id!r} resolveu para parágrafo sem runs"
            continue
        documento_docx.add_comment(
            paragrafo.runs, text=_texto_do_comentario(c), author=autor, initials=iniciais
        )
        resultado.comentarios_aplicados.append(c.comment_id)

    documento_docx.save(caminho_saida)
    return resultado
