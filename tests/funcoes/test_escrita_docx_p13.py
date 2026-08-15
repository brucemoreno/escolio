"""Testes de `escolio.funcoes.escrita_docx_p13` — primeira escrita real de
comentário do Word a partir de `P13Comment` [sessão de 2026-08-14].
"""

import docx
import pytest

from escolio.comentarios.comentario import P13Comment
from escolio.comentarios.vocabulario import P13CommentStatus
from escolio.funcoes.escrita_docx_p13 import escrever_comentarios_no_docx
from escolio.ingestao.parser_docx import parse_docx


def _comentario(unit_id: str, document_id: str, comment_id: str = "CMT-TESTE-0001") -> P13Comment:
    return P13Comment(
        comment_id=comment_id,
        document_id=document_id,
        document_version="1.0.0",
        module_id="P13",
        unit_id=unit_id,
        anchor_start="0",
        anchor_end="10",
        anchor_text_hash="sha256:teste",
        comment_type="COMENTARIO_FACTUAL",
        priority="PRIORIDADE_MEDIA",
        severity="MODERADA",
        problem="Problema de teste.",
        evidence="Evidência de teste.",
        impact="Impacto de teste.",
        recommended_action="Ação recomendada de teste.",
        intervention_level="INT-04",
        authority_required="USUARIO_PROPONENTE",
        gate="GATE_DE_VALIDACAO_FINAL",
        source_status="VERIFICADA",
        voice_impact="NENHUM",
        privacy_classification="PUBLIC",
        reversible=True,
        status=P13CommentStatus.DRAFT,
    )


def _docx_sintetico(tmp_path) -> str:
    caminho = str(tmp_path / "sintetico.docx")
    d = docx.Document()
    d.add_paragraph("Primeiro parágrafo do corpo, usado para ancorar comentário.")
    d.add_paragraph("Segundo parágrafo do corpo, sem comentário.")
    d.save(caminho)
    return caminho


class TestEscreverComentariosNoDocx:
    def test_recusa_sobrescrever_o_original(self, tmp_path):
        caminho = _docx_sintetico(tmp_path)
        with pytest.raises(ValueError, match="não pode ser igual"):
            escrever_comentarios_no_docx(caminho, [], caminho)

    def test_comentario_real_e_gravado_ancorado_e_original_preservado(self, tmp_path):
        caminho_entrada = _docx_sintetico(tmp_path)
        caminho_saida = str(tmp_path / "com_comentarios.docx")
        documento = parse_docx(caminho_entrada)
        alvo = documento.paragrafos[0]

        resultado = escrever_comentarios_no_docx(
            caminho_entrada, [_comentario(alvo.unit_id, documento.hash_documento)], caminho_saida
        )

        assert resultado.comentarios_aplicados == ["CMT-TESTE-0001"]
        assert resultado.nao_aplicados == {}

        salvo = docx.Document(caminho_saida)
        comentarios_no_arquivo = list(salvo.comments)
        assert len(comentarios_no_arquivo) == 1
        assert "Problema de teste" in comentarios_no_arquivo[0].paragraphs[0].text

        original = docx.Document(caminho_entrada)
        assert list(original.comments) == []

    def test_unit_id_desconhecido_vai_para_nao_aplicados_sem_falhar(self, tmp_path):
        caminho_entrada = _docx_sintetico(tmp_path)
        caminho_saida = str(tmp_path / "com_comentarios.docx")
        documento = parse_docx(caminho_entrada)

        resultado = escrever_comentarios_no_docx(
            caminho_entrada,
            [_comentario("UNI-PAR-NAO-EXISTE", documento.hash_documento)],
            caminho_saida,
        )

        assert resultado.comentarios_aplicados == []
        assert "CMT-TESTE-0001" in resultado.nao_aplicados
        assert "não resolve" in resultado.nao_aplicados["CMT-TESTE-0001"]

    def test_multiplos_comentarios_multiplos_paragrafos(self, tmp_path):
        caminho_entrada = _docx_sintetico(tmp_path)
        caminho_saida = str(tmp_path / "com_comentarios.docx")
        documento = parse_docx(caminho_entrada)
        c1 = _comentario(documento.paragrafos[0].unit_id, documento.hash_documento, "CMT-A")
        c2 = _comentario(documento.paragrafos[1].unit_id, documento.hash_documento, "CMT-B")

        resultado = escrever_comentarios_no_docx(caminho_entrada, [c1, c2], caminho_saida)

        assert set(resultado.comentarios_aplicados) == {"CMT-A", "CMT-B"}
        salvo = docx.Document(caminho_saida)
        assert len(list(salvo.comments)) == 2
